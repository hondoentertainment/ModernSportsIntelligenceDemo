#!/usr/bin/env python3
"""Generate the MSI Feature Guide as a Word document."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color_hex,
    })
    shading.append(shd)


def styled_table(doc, headers, rows, header_color='1B2A4A'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(9)
        set_cell_shading(cell, header_color)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_shading(cell, 'F2F6FA')
    return table


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def bold_para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    return p


def main():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.54)
        s.bottom_margin = Cm(2.54)
        s.left_margin = Cm(2.54)
        s.right_margin = Cm(2.54)

    # Styles
    for level, size, color in [
        ('Heading 1', 22, (0x1B, 0x2A, 0x4A)),
        ('Heading 2', 16, (0x1B, 0x2A, 0x4A)),
        ('Heading 3', 13, (0x2D, 0x4A, 0x7A)),
    ]:
        st = doc.styles[level]
        st.font.color.rgb = RGBColor(*color)
        st.font.size = Pt(size)

    # ---- TITLE PAGE ----
    for _ in range(6):
        doc.add_paragraph('')
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('Modern Sports Intelligence')
    r.bold = True; r.font.size = Pt(36); r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t2.add_run('Feature Guide')
    r.font.size = Pt(22); r.font.color.rgb = RGBColor(0x5A, 0x6A, 0x8A)

    doc.add_paragraph('')
    v = doc.add_paragraph()
    v.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = v.add_run(f'Version 10.0  |  {datetime.date.today().strftime("%B %d, %Y")}')
    r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    tl = doc.add_paragraph()
    tl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tl.add_run('A comprehensive guide to every feature in the platform')
    r.italic = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x66, 0x77, 0x88)

    doc.add_page_break()

    # ---- TABLE OF CONTENTS ----
    doc.add_heading('Table of Contents', level=1)
    toc = [
        '1. Getting Started',
        '2. Dashboard & Portfolio Overview',
        '3. Collection Management',
        '4. Market Intelligence',
        '5. AI-Powered Features',
        '6. Trading & Execution Tools',
        '7. Analytics & Charting',
        '8. Grading & Condition Tools',
        '9. Financial & Tax Tools',
        '10. Social & Community',
        '11. League-Specific Hubs',
        '12. Enterprise & Team Features',
        '13. Advanced Quantitative Tools',
        '14. Industry-First Innovations',
        '15. Mobile & Accessibility',
        '16. Settings & Account Management',
    ]
    for item in toc:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)
    doc.add_page_break()

    # =====================================================================
    # 1. GETTING STARTED
    # =====================================================================
    doc.add_heading('1. Getting Started', level=1)

    doc.add_heading('1.1 Creating Your Account', level=2)
    doc.add_paragraph(
        'Navigate to the login page to create a new account using email and password. '
        'MSI uses Supabase authentication for secure account management. A demo mode is '
        'available for exploring the platform without creating an account.'
    )
    bold_para(doc, 'How to access:')
    bullets(doc, [
        'Click "Sign Up" on the login screen (/login)',
        'Enter your email and create a password',
        'Verify your email address via the confirmation link',
        'Or click "Try Demo" to explore without an account',
    ])

    doc.add_heading('1.2 Importing Your Collection', level=2)
    doc.add_paragraph(
        'Bring your existing collection into MSI through multiple import methods.'
    )
    bold_para(doc, 'Import options:')
    bullets(doc, [
        'CSV Import: Upload a spreadsheet with your card data (player, year, set, value, grade)',
        'eBay Import: Pull your purchase history directly from eBay',
        'PSA/BGS Lookup: Enter certification numbers to auto-populate card details',
        'OCR Scanner: Use your phone camera to scan cards for automatic identification',
        'Manual Entry: Add cards one at a time with full detail forms',
    ])

    doc.add_heading('1.3 Guided Tour', level=2)
    doc.add_paragraph(
        'First-time users are offered an interactive guided tour that walks through the '
        'key features of the platform. The tour highlights the dashboard, collection management, '
        'watchlist, and AI-powered tools.'
    )

    doc.add_heading('1.4 Navigation', level=2)
    doc.add_paragraph(
        'The sidebar provides access to all major sections. Use the Command Palette (Ctrl+K / Cmd+K) '
        'for instant access to any feature, card, or action from anywhere in the app.'
    )

    doc.add_page_break()

    # =====================================================================
    # 2. DASHBOARD
    # =====================================================================
    doc.add_heading('2. Dashboard & Portfolio Overview', level=1)

    doc.add_heading('2.1 League Intelligence Dashboard', level=2)
    doc.add_paragraph('Route: / (Home)')
    doc.add_paragraph(
        'Your command center for portfolio monitoring. The dashboard displays real-time portfolio '
        'health at a glance.'
    )
    bold_para(doc, 'What you will see:')
    bullets(doc, [
        'Net Asset Value (NAV): Total current value of your collection with daily/weekly/monthly change',
        'Profit & Loss: Unrealized gains and losses across all holdings with percentage returns',
        'Top Movers: Cards with the biggest price changes today',
        'Recent Activity: Latest transactions, price alerts, and sync updates',
        'Market Sentiment: Bullish/bearish indicators across leagues',
        'Quick Stats: Total cards, graded count, average ROI, diversification score',
    ])

    doc.add_heading('2.2 Market Ticker', level=2)
    doc.add_paragraph(
        'A real-time scrolling ticker in the page header shows live price movements for '
        'trending cards across all leagues. Click any ticker item to view full details.'
    )

    doc.add_heading('2.3 Morning Briefing', level=2)
    doc.add_paragraph(
        'An AI-generated daily summary widget that highlights overnight market movements, '
        'triggered watchlist alerts, portfolio changes, and recommended actions for the day.'
    )

    doc.add_page_break()

    # =====================================================================
    # 3. COLLECTION MANAGEMENT
    # =====================================================================
    doc.add_heading('3. Collection Management', level=1)

    doc.add_heading('3.1 Collection Inventory', level=2)
    doc.add_paragraph('Route: /collection')
    doc.add_paragraph(
        'The heart of MSI - a high-performance inventory system for managing your entire card portfolio.'
    )
    bold_para(doc, 'Features:')
    bullets(doc, [
        'Virtualized Grid: Handles 10,000+ cards smoothly with grid, list, and compact view modes',
        'Search & Filter: Filter by player, year, sport, manufacturer, grade, status, and custom tags',
        'Sort Options: Sort by value, ROI, player name, date added, grade, or scarcity',
        'Card Details: Click any card to view full details including price history, comps, and analytics',
        'Image Management: Upload card images with lazy loading and full-screen lightbox viewer',
    ])

    doc.add_heading('3.2 Adding & Editing Cards', level=2)
    doc.add_paragraph(
        'Add cards manually with comprehensive detail forms or use bulk import tools. '
        'Each card tracks: player, year, manufacturer, set, card number, sport/league, '
        'condition, grading info, purchase price/date, current value, and notes.'
    )

    doc.add_heading('3.3 Bulk Operations', level=2)
    doc.add_paragraph(
        'Select multiple cards using checkboxes for batch actions.'
    )
    bold_para(doc, 'Available bulk actions:')
    bullets(doc, [
        'Bulk Edit: Update fields across multiple cards simultaneously',
        'Bulk Delete: Remove multiple cards at once',
        'Bulk Export: Export selected cards to CSV or PDF',
        'Bulk Status Change: Move cards between owned/vaulted/pending',
    ])

    doc.add_heading('3.4 Sold Vault', level=2)
    doc.add_paragraph(
        'When you sell a card, it moves to the Sold Vault rather than being deleted. '
        'This preserves your complete transaction history and allows tracking of realized '
        'gains/losses for performance analysis and tax reporting.'
    )

    doc.add_heading('3.5 Card Status Tracking', level=2)
    styled_table(doc,
        ['Status', 'Description'],
        [
            ['Owned', 'Card is in your active collection'],
            ['Sold', 'Card has been sold (moves to Sold Vault with realized P/L)'],
            ['Pending', 'Card is in transit or awaiting delivery'],
            ['Vaulted', 'Card is stored with a vault provider for safekeeping'],
        ]
    )

    doc.add_heading('3.6 Set Registry', level=2)
    doc.add_paragraph(
        'Track progress toward completing full card sets. The registry shows which cards '
        'you own, which are missing, and the market cost to complete each set. Visual '
        'progress bars and completion percentages help prioritize acquisitions.'
    )

    doc.add_heading('3.7 Digital Display Case', level=2)
    doc.add_paragraph(
        'Create themed showcases of your best cards with shareable URLs. Choose from '
        'multiple display themes and arrange cards to highlight your collection\'s highlights.'
    )

    doc.add_page_break()

    # =====================================================================
    # 4. MARKET INTELLIGENCE
    # =====================================================================
    doc.add_heading('4. Market Intelligence', level=1)

    doc.add_heading('4.1 Market Pulse', level=2)
    doc.add_paragraph(
        'Real-time sentiment analysis across the sports card market. The Market Pulse engine '
        'aggregates signals from social media, marketplace activity, and news to produce '
        'sentiment scores (bullish/bearish/neutral) broken down by league and category.'
    )
    bold_para(doc, 'Key indicators:')
    bullets(doc, [
        'Alpha Correlation: Links market sentiment shifts to your portfolio positions',
        'Lagging Alpha Alerts: Identifies when sentiment has shifted but prices haven\'t yet - potential early trade signals',
        'Hype Feed: Real-time social signal aggregation showing trending cards and narratives',
    ])

    doc.add_heading('4.2 Market Trends', level=2)
    doc.add_paragraph('Route: /trends')
    doc.add_paragraph(
        'Multi-dimensional trend analysis across time periods, categories, players, and leagues. '
        'View trending cards, rising players, and market-wide patterns.'
    )

    doc.add_heading('4.3 Scarcity Intelligence', level=2)
    doc.add_paragraph(
        'Population reporting and scarcity analysis powered by PSA and BGS population data.'
    )
    bold_para(doc, 'Features:')
    bullets(doc, [
        'Pop Count Tracking: Current population counts for each card/grade combination',
        'Scarcity Badges: Visual indicators showing relative scarcity (Pop 1, Low Pop, etc.)',
        'Pop 1 Alerts: Notifications when you own or can acquire one-of-one graded cards',
        'Scarcity Index: Computed score combining pop count, demand signals, and market activity',
    ])

    doc.add_heading('4.4 Anomaly Detection', level=2)
    doc.add_paragraph(
        'Statistical anomaly detection identifies unusual price movements and market activity. '
        'When a card\'s price deviates significantly from expected patterns, the system flags it '
        'for review - potentially signaling buying or selling opportunities.'
    )

    doc.add_heading('4.5 Macro-Sentinel', level=2)
    doc.add_paragraph(
        'Early warning system for macro events that could impact the card market. Monitors '
        'economic indicators, cultural events, media releases, and sports milestones that '
        'historically correlate with market movements.'
    )

    doc.add_heading('4.6 Deal Finder & Arbitrage', level=2)
    doc.add_paragraph(
        'Scans multiple marketplaces to identify underpriced cards and arbitrage opportunities. '
        'Each opportunity is scored based on discount percentage, liquidity, and confidence level.'
    )

    doc.add_page_break()

    # =====================================================================
    # 5. AI-POWERED FEATURES
    # =====================================================================
    doc.add_heading('5. AI-Powered Features', level=1)
    doc.add_paragraph(
        'MSI integrates Google Gemini for advanced AI capabilities throughout the platform.'
    )

    doc.add_heading('5.1 AI Discovery Engine', level=2)
    doc.add_paragraph(
        'Automated card valuation and prospect discovery. The AI analyzes market data, '
        'player trajectories, and historical patterns to surface investment opportunities '
        'and provide confidence-scored valuations.'
    )

    doc.add_heading('5.2 Deep Search Intelligence', level=2)
    doc.add_paragraph('Route: /deep-search')
    doc.add_paragraph(
        'Go beyond keyword search with natural language queries.'
    )
    bold_para(doc, 'Example searches:')
    bullets(doc, [
        '"Rookie cards of players having breakout seasons under $50"',
        '"Junk wax era cards with future Hall of Fame potential"',
        '"High-grade vintage cards with low population counts"',
        '"Cards similar to my top performers"',
    ])

    doc.add_heading('5.3 AI Negotiation Arena', level=2)
    doc.add_paragraph(
        'Practice and conduct negotiations with AI-powered seller agents. The system tracks '
        'negotiation sentiment, suggests counter-offers based on market data, and provides '
        'real-time deal analysis comparing offers to fair market value.'
    )

    doc.add_heading('5.4 Smart Collection Advisor', level=2)
    doc.add_paragraph(
        'AI-generated buy/sell recommendations based on your portfolio composition, market '
        'trends, risk tolerance, and investment goals. Recommendations include rationale '
        'and confidence scoring.'
    )

    doc.add_heading('5.5 AI Predictive Grading', level=2)
    doc.add_paragraph(
        'Computer vision analysis predicts professional grades before submission. Upload '
        'card photos to receive estimated grades with confidence intervals, helping optimize '
        'grading submissions for maximum ROI.'
    )

    doc.add_heading('5.6 Counterfeit Detection', level=2)
    doc.add_paragraph(
        'AI-powered authenticity analysis scores cards for counterfeit risk. Includes '
        'certificate verification, image analysis, and pattern matching against known '
        'forgery databases.'
    )

    doc.add_heading('5.7 Card Genome Sequencer', level=2)
    doc.add_paragraph('Route: /card-genome-sequencer')
    doc.add_paragraph(
        'Creates a unique 18-gene DNA fingerprint for every card across Physical, Visual, '
        'Chemical, and Rarity categories. Enables sophisticated similarity matching, '
        'collection gap analysis, and pattern-based discovery.'
    )

    doc.add_heading('5.8 Narrative Arc Engine', level=2)
    doc.add_paragraph('Route: /narrative-arc-engine')
    doc.add_paragraph(
        'Evaluates player career narratives (comeback stories, dynasty builders, tragic figures) '
        'and their impact on card values. Some of the most valuable cards gain premiums from '
        'compelling stories, not just statistics.'
    )

    doc.add_page_break()

    # =====================================================================
    # 6. TRADING & EXECUTION
    # =====================================================================
    doc.add_heading('6. Trading & Execution Tools', level=1)

    doc.add_heading('6.1 Watchlist & Price Alerts', level=2)
    doc.add_paragraph('Route: /favorites')
    doc.add_paragraph(
        'Track cards you want to acquire with target prices. Set alerts to be notified '
        'when market prices drop below your target. Priority levels (high/medium/low) '
        'help organize acquisition planning.'
    )

    doc.add_heading('6.2 Auction Sniper', level=2)
    doc.add_paragraph('Route: /auction-sniper')
    doc.add_paragraph(
        'Automated bid timing system for marketplace auctions. Analyzes auction patterns '
        'and places optimally-timed bids to maximize win rate at minimum price. Supports '
        'multiple marketplace integrations.'
    )

    doc.add_heading('6.3 Trade Block & Offers', level=2)
    doc.add_paragraph(
        'Organize cards available for trade and evaluate incoming trade packages. The system '
        'provides fair-value analysis for trade offers, showing whether each side is getting '
        'comparable value.'
    )

    doc.add_heading('6.4 Consignment Tracker', level=2)
    doc.add_paragraph(
        'Track cards sent to consignment houses (PWCC, Goldin, Heritage, etc.). Monitors '
        'listing status, sale results, and fee reconciliation. Local-first architecture '
        'ensures your data is always accessible.'
    )

    doc.add_heading('6.5 Rules Engine', level=2)
    doc.add_paragraph(
        'Set conditional trading rules that execute automatically when conditions are met. '
        'Example rules: "Sell if price exceeds 2x purchase price" or "Buy if price drops '
        'below $X." Includes backtesting against historical data.'
    )

    doc.add_heading('6.6 Private Deal Room', level=2)
    doc.add_paragraph('Route: /private-deal-room-agent')
    doc.add_paragraph(
        'Encrypted negotiation rooms for high-value private transactions. AI agent support '
        'provides real-time market data, valuation, and counter-offer suggestions during '
        'negotiations.'
    )

    doc.add_heading('6.7 eBay Listing Generator', level=2)
    doc.add_paragraph('Route: /ebay-listing-generator')
    doc.add_paragraph(
        'Generates SEO-optimized eBay listings with professional templates, keyword '
        'optimization, and comp-based pricing suggestions. Streamlines the selling process '
        'for maximum visibility and returns.'
    )

    doc.add_heading('6.8 Institutional Liquidity Pool', level=2)
    doc.add_paragraph(
        'Instant-liquidity marketplace where high-value cards can be sold immediately '
        'to institutional buyers at a known discount. Provides guaranteed exit liquidity '
        'when you need to sell quickly.'
    )

    doc.add_page_break()

    # =====================================================================
    # 7. ANALYTICS & CHARTING
    # =====================================================================
    doc.add_heading('7. Analytics & Charting', level=1)

    doc.add_heading('7.1 Price History Charts', level=2)
    doc.add_paragraph(
        'Detailed time-series charts for any card with technical overlays.'
    )
    bold_para(doc, 'Available overlays:')
    bullets(doc, [
        'Simple Moving Average (SMA) - 7-day, 30-day, 90-day',
        'Bollinger Bands for volatility analysis',
        'Volume indicators showing market activity',
        'Statistical annotations (mean, median, standard deviation)',
    ])

    doc.add_heading('7.2 Technical Analysis Suite', level=2)
    doc.add_paragraph(
        'Full charting toolkit adapted for sports card price data.'
    )
    bullets(doc, [
        'RSI (Relative Strength Index) for overbought/oversold signals',
        'MACD (Moving Average Convergence Divergence) for trend identification',
        'Fibonacci retracements for support/resistance levels',
        'Volume profile analysis for price level interest',
    ])

    doc.add_heading('7.3 Portfolio Audit Engine', level=2)
    doc.add_paragraph('Route: /audit')
    doc.add_paragraph(
        'Comprehensive portfolio health analysis covering ROI by category, diversification '
        'scoring, risk assessment, concentration warnings, and actionable improvement recommendations.'
    )

    doc.add_heading('7.4 Comparative Analysis', level=2)
    doc.add_paragraph('Route: /compare')
    doc.add_paragraph(
        'Side-by-side comparison of cards or players across multiple attributes including '
        'price performance, scarcity, liquidity, grade premiums, and investment potential.'
    )

    doc.add_heading('7.5 Analyst War Room', level=2)
    doc.add_paragraph('Route: /war-room')
    doc.add_paragraph(
        'Professional multi-panel workspace for simultaneous monitoring. Arrange multiple '
        'data views - charts, tables, feeds, and analytics - in a customizable layout '
        'for comprehensive market monitoring.'
    )

    doc.add_heading('7.6 Cross-Correlation Engine', level=2)
    doc.add_paragraph(
        'Analyze correlations between card categories, leagues, and even traditional financial '
        'assets (S&P 500, crypto, bonds). Helps build diversified portfolios and identify '
        'hedging opportunities.'
    )

    doc.add_heading('7.7 Liquidity Intelligence', level=2)
    doc.add_paragraph(
        'Institutional-grade liquidity analysis featuring order book simulation, volume '
        'velocity tracking, and market depth heatmaps. Understand how quickly and at what '
        'price you can exit any position.'
    )

    doc.add_heading('7.8 What-If Simulator', level=2)
    doc.add_paragraph('Route: /what-if-simulator')
    doc.add_paragraph(
        'Model the impact of potential trades on your portfolio. Simulate buying, selling, '
        'or trading cards to see effects on NAV, diversification, risk metrics, and ROI '
        'before committing to any transaction.'
    )

    doc.add_heading('7.9 Portfolio Time Machine', level=2)
    doc.add_paragraph(
        'Replay historical portfolio states to analyze past decisions. Includes a decision '
        'journal for annotating why you made specific trades, enabling learning from '
        'both successes and mistakes.'
    )

    doc.add_page_break()

    # =====================================================================
    # 8. GRADING & CONDITION
    # =====================================================================
    doc.add_heading('8. Grading & Condition Tools', level=1)

    doc.add_heading('8.1 Grading Batch Planner', level=2)
    doc.add_paragraph('Route: /grading-batch-planner')
    doc.add_paragraph(
        'Optimizes which cards to submit for professional grading and how to batch them. '
        'Calculates expected ROI for each card at various grade outcomes, recommends '
        'submission tiers (economy, regular, express), and groups cards into cost-effective batches.'
    )

    doc.add_heading('8.2 Visual Audit Simulation', level=2)
    doc.add_paragraph(
        'Computer vision grading simulation that analyzes card photos and provides estimated '
        'grades with ROI overlays. Shows the expected value increase at each potential grade '
        'to help decide whether professional grading is worth the cost.'
    )

    doc.add_heading('8.3 Grade Premium Analytics', level=2)
    doc.add_paragraph(
        'Quantifies the value premium between adjacent grades. See exactly how much more '
        'a PSA 10 commands over a PSA 9 for any card, and track how grade premiums shift '
        'over time. Also calculates crossover ROI (e.g., cracking a BGS 9.5 to submit to PSA).'
    )

    doc.add_heading('8.4 Pop Report Growth Tracker', level=2)
    doc.add_paragraph(
        'Monitors how grading populations change over time. Rising pop counts can erode '
        'scarcity premiums - this tool tracks trends and alerts you when populations are '
        'growing faster than demand for cards you own.'
    )

    doc.add_heading('8.5 Restoration Simulator', level=2)
    doc.add_paragraph('Route: /restoration-simulator')
    doc.add_paragraph(
        'Models the ROI of cracking graded cards and resubmitting for a potentially higher '
        'grade. Factors in cracking risk, resubmission costs, and probability distributions '
        'for grade outcomes.'
    )

    doc.add_heading('8.6 Forensics Lab', level=2)
    doc.add_paragraph('Route: /forensics-lab')
    doc.add_paragraph(
        '10-layer counterfeit detection system using AI image analysis, certificate verification, '
        'materials analysis, font matching, and comparison against known authentic examples.'
    )

    doc.add_page_break()

    # =====================================================================
    # 9. FINANCIAL & TAX
    # =====================================================================
    doc.add_heading('9. Financial & Tax Tools', level=1)

    doc.add_heading('9.1 Fiscal Intelligence', level=2)
    doc.add_paragraph(
        'Comprehensive tax management for card trading. Tracks cost basis using FIFO, LIFO, '
        'or specific identification methods. Generates Schedule D exports for tax filing.'
    )

    doc.add_heading('9.2 Tax-Loss Harvesting', level=2)
    doc.add_paragraph(
        'Identifies strategic loss-selling opportunities to offset capital gains. Automatically '
        'accounts for wash-sale rules to ensure tax compliance while optimizing your tax position.'
    )

    doc.add_heading('9.3 Break-Even Calculator', level=2)
    doc.add_paragraph(
        'Calculates the minimum sale price needed to break even after marketplace fees '
        '(eBay, PWCC, Goldin), shipping costs, and applicable taxes. Prevents selling at a '
        'hidden loss.'
    )

    doc.add_heading('9.4 Insurance Valuation Report', level=2)
    doc.add_paragraph(
        'Generates replacement-cost estimates for insurance documentation. Provides card-by-card '
        'valuations with market data sources, total collection value, and formatted reports '
        'suitable for insurance claims.'
    )

    doc.add_heading('9.5 Insurance Manager', level=2)
    doc.add_paragraph(
        'Track insurance policies, premium payments, and coverage gaps. Alerts you when '
        'collection value exceeds current coverage limits.'
    )

    doc.add_heading('9.6 Collector Audit Dossier', level=2)
    doc.add_paragraph('Route: /audit-dossier')
    doc.add_paragraph(
        'One-click export combining tax documentation, insurance valuations, and exit planning '
        'into a single professional report. Ideal for year-end financial review or estate planning.'
    )

    doc.add_heading('9.7 Multi-Currency Support', level=2)
    doc.add_paragraph(
        'Exchange rate conversion, cross-border price comparison, and import duty calculations '
        'for international collectors buying from global marketplaces.'
    )

    doc.add_heading('9.8 Wax Break ROI Tracker', level=2)
    doc.add_paragraph('Route: /wax-break-roi-tracker')
    doc.add_paragraph(
        'Log sealed wax product purchases (boxes, cases, packs) and track the cards pulled. '
        'Calculates true ROI on wax investments including all hits and bulk, helping determine '
        'whether sealed product is a good investment.'
    )

    doc.add_page_break()

    # =====================================================================
    # 10. SOCIAL & COMMUNITY
    # =====================================================================
    doc.add_heading('10. Social & Community', level=1)

    doc.add_heading('10.1 Social Alpha Elite', level=2)
    doc.add_paragraph(
        'Live hype feed aggregating social signals from across the hobby. Tracks trending '
        'cards, viral moments, and community narratives. Institutional collector tiers '
        'highlight moves by top performers.'
    )

    doc.add_heading('10.2 Social Network', level=2)
    doc.add_paragraph(
        'Full social platform with follower/following relationships, direct messaging, '
        'trade matching based on complementary collections, and activity feeds showing '
        'what top collectors are buying and selling.'
    )

    doc.add_heading('10.3 Community Benchmarking', level=2)
    doc.add_paragraph('Route: /leaderboard')
    doc.add_paragraph(
        'Performance leaderboards ranking collectors by ROI, total collection value, '
        'alpha generation, and collecting milestones. See how you stack up against the community.'
    )

    doc.add_heading('10.4 Achievements System', level=2)
    doc.add_paragraph(
        '30+ badges across 5 tiers with level progression. Earn achievements for collection '
        'milestones (first PSA 10, completing a set), trading milestones (profitable trades), '
        'and engagement milestones (daily logins, community contributions).'
    )

    doc.add_heading('10.5 Collection Sharing', level=2)
    doc.add_paragraph(
        'Generate embeddable HTML widgets to showcase your collection on personal websites, '
        'forums, or social media profiles. Public portfolio pages with shareable URLs.'
    )

    doc.add_heading('10.6 Copy Trading', level=2)
    doc.add_paragraph(
        'Follow top-performing collectors and optionally mirror their portfolio moves. '
        'See what the best performers are buying and selling with fractional ownership concepts.'
    )

    doc.add_heading('10.7 Counterparty Trust Graph', level=2)
    doc.add_paragraph('Route: /counterparty-trust-graph')
    doc.add_paragraph(
        'Seller trust scoring derived from completed deals, referrals, and dispute outcomes. '
        'Before transacting with someone, check their trust profile to reduce fraud risk.'
    )

    doc.add_page_break()

    # =====================================================================
    # 11. LEAGUE HUBS
    # =====================================================================
    doc.add_heading('11. League-Specific Hubs', level=1)
    doc.add_paragraph(
        'Each major sport has a dedicated hub with sport-specific intelligence and analytics.'
    )

    for league, route, desc, features in [
        ('MLB', '/mlb-stats', 'Baseball-specific analytics with live MLB Stats API integration.',
         ['Live player stats and performance data', 'PressBox Hub for real-time game data',
          'Prospect pipeline from MiLB to majors', 'Seasonal pricing patterns (Spring Training, playoffs, HOF voting)']),
        ('NFL', '/nfl-hub', 'Football-specific intelligence for the largest sports card market.',
         ['Draft capital tracking and rookie impact modeling', 'Injury report integration with card value impact',
          'Position-specific valuation (QB premium analysis)', 'Super Bowl and playoff performance premiums']),
        ('NBA', '/nba-hub', 'Basketball card intelligence with trade and free agency modeling.',
         ['Trade deadline and free agency impact analysis', 'Rookie scale contract correlation',
          'All-Star and playoff performance premium tracking', 'International prospect pipeline monitoring']),
        ('NHL', '/nhl-hub', 'Hockey card analytics with Young Guns and award tracking.',
         ['Young Guns and rookie card set tracking', 'Playoff performance value correlations',
          'Award race tracking (Hart, Calder, etc.)', 'International scouting integration']),
        ('Soccer', '/soccer-hub', 'Global football card intelligence across multiple leagues.',
         ['Transfer window impact modeling', 'World Cup and Euro championship price effects',
          'Multi-league coverage (Premier League, La Liga, Serie A, Bundesliga)',
          'Youth academy prospect tracking']),
    ]:
        doc.add_heading(f'{league} Hub ({route})', level=2)
        doc.add_paragraph(desc)
        bullets(doc, features)

    doc.add_heading('11.5 Team Analytics', level=2)
    doc.add_paragraph('Route: /teams')
    doc.add_paragraph(
        'Team-level analytics including offense/defense ratings, form indicators, and how '
        'team performance impacts individual player card values.'
    )

    doc.add_heading('11.6 Game Intelligence', level=2)
    doc.add_paragraph('Route: /games')
    doc.add_paragraph(
        'Live game tracking with player impact scoring. See real-time performance data '
        'during games and understand how in-game events affect card values.'
    )

    doc.add_page_break()

    # =====================================================================
    # 12. ENTERPRISE
    # =====================================================================
    doc.add_heading('12. Enterprise & Team Features', level=1)
    doc.add_paragraph(
        'Collaborative tools for collector groups, clubs, dealers, and investment funds.'
    )

    categories = [
        ('Team Collaboration', [
            ('Shared Watchlists', 'Real-time synchronized watchlists with shared notes and alerts across team members'),
            ('Shared Dashboards', 'Aggregated portfolio views across multiple collectors with combined analytics'),
            ('Draft Night War Room', 'Synchronized draft tracking with real-time card market impact across team'),
            ('Card Show Squad', 'Coordinate group attendance at shows: booth sharing, target lists, budget planning'),
        ]),
        ('Group Trading', [
            ('Group Breaks', 'Live group break participation with automated card distribution and hit tracking'),
            ('Group Buy Coop', 'Cooperative bulk purchasing with cost-sharing and distribution models'),
            ('Swap Meet', 'Peer-to-peer trading marketplace within trusted collector networks'),
            ('Dispute Arbitration', 'Transaction dispute resolution with evidence management and mediator support'),
        ]),
        ('Club & Organization', [
            ('Club Management', 'Guild/club operations: membership, dues, shared resources, event planning'),
            ('Collective Grading', 'Shared grading submissions for cost-efficient batch grading across members'),
            ('Collaborative Set Registry', 'Team-based set completion where multiple collectors contribute cards'),
            ('Mentorship Exchange', 'Peer learning platform connecting experienced collectors with newcomers'),
        ]),
        ('Enterprise Administration', [
            ('RBAC Teams', 'Role-based access control with granular permissions for enterprise accounts'),
            ('White-Label API', 'Partner integration API for embedding MSI into third-party platforms'),
            ('Dealer Inventory', 'Dealer-specific tools: wholesale pricing, consignment, customer management'),
            ('Fund Reporting', 'Institutional-grade NAV calculations and investor communications'),
            ('Audit Trail', 'Compliance-grade logging for all portfolio actions and transactions'),
        ]),
    ]

    for cat_name, features in categories:
        doc.add_heading(cat_name, level=2)
        for fname, fdesc in features:
            doc.add_heading(fname, level=3)
            doc.add_paragraph(fdesc)

    doc.add_page_break()

    # =====================================================================
    # 13. ADVANCED QUANTITATIVE
    # =====================================================================
    doc.add_heading('13. Advanced Quantitative Tools', level=1)
    doc.add_paragraph(
        'Institutional-grade quantitative analysis tools adapted from traditional finance.'
    )

    quant_features = [
        ('Monte Carlo Stress Testing', 'Run 1000+ simulation paths on your portfolio to model potential outcomes. Calculate Value at Risk (VaR), maximum drawdown, and tail-risk scenarios. Understand how your portfolio might perform under various market conditions.'),
        ('Cross-Asset Correlation', 'Correlate card price movements with the S&P 500, Bitcoin, gold, real estate, and treasury yields. Identify macro trends affecting card values and build truly diversified portfolios.'),
        ('Portfolio Scenario Theater (/portfolio-scenario-theater)', 'Persist and compare multiple Monte Carlo scenarios with hedging guidance. Save scenarios to track how predictions evolve versus actual outcomes.'),
        ('Chaos Theory Simulator (/chaos-theory-simulator)', 'Butterfly effect modeling for sports card markets. Simulate how small events (a player interview, a social media post) can cascade into major market movements through network effects.'),
        ('Contagion Mapper (/contagion-mapper)', 'Systemic risk propagation modeling. Visualize how a crash in one card category (e.g., basketball vintage) might cascade through the broader market.'),
        ('Rebalancing Alerts', 'Set target allocation weights by league, category, or player. Receive alerts when portfolio drift exceeds your thresholds, with suggested trades to rebalance.'),
        ('Seasonal Strategy Engine', 'Data-driven buy/sell timing recommendations based on historical sport season cycles. Know when prices typically peak and trough for each sport and card category.'),
    ]

    for title, desc in quant_features:
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)

    doc.add_page_break()

    # =====================================================================
    # 14. INDUSTRY-FIRST
    # =====================================================================
    doc.add_heading('14. Industry-First Innovations', level=1)
    doc.add_paragraph(
        'Capabilities that do not exist in any competing platform - MSI exclusives.'
    )

    innovations = [
        ('Injury Oracle (/injury-oracle)', 'Forecasts player injury probability and models the downstream impact on card values. Helps you manage risk in portfolios concentrated on injury-prone players.'),
        ('Market Maker Arena (/market-maker-arena)', 'Design, backtest, and deploy competing AI trading bots. Watch multiple strategies compete in real-time with live profit/loss tracking.'),
        ('Fractional Vault (/fractional-vault)', 'Tokenized fractional ownership of high-value cards. Multiple collectors can co-own a $50,000 card with governance mechanisms for decisions like selling or grading.'),
        ('Card Decay Predictor (/card-decay-predictor)', 'Physics-based degradation modeling predicting how card condition deteriorates over time. Factors in storage conditions (temperature, humidity, UV) to estimate half-life.'),
        ('Print Run Decoder (/print-run-decoder)', 'Reverse-engineers hidden print run numbers from population reports, marketplace data, and statistical modeling. Reveals the true scarcity of cards whose print runs were never disclosed.'),
        ('Card Yield Farming (/card-yield-farming)', 'Generate returns from your collection beyond price appreciation: lending for display, licensing images, staking in insurance pools, and content creation royalties.'),
        ('Temporal Arbitrage Radar (/temporal-arbitrage-radar)', 'Detects cyclical and seasonal mispricing patterns. Identifies when specific cards are historically underpriced relative to upcoming catalysts.'),
        ('Collection DNA Mixer (/collection-dna-mixer)', 'Genetic algorithm portfolio optimization that evolves collection compositions toward optimal risk/return profiles through simulated natural selection.'),
        ('IoT Condition Guardian', 'Smart sensor integration for monitoring storage conditions. Connect temperature, humidity, and UV sensors to receive alerts when conditions threaten card integrity.'),
        ('Dopamine Cycle Tracker', 'Behavioral finance self-awareness tool. Tracks your buying and selling patterns to identify emotional decision-making, FOMO purchases, and panic selling.'),
        ('HOF Probability Engine', 'Models career trajectories to estimate Hall of Fame probability. Long-term card values are heavily influenced by HOF status - this tool helps assess that likelihood.'),
        ('AR Vault Walkthrough', 'Immersive 3D augmented reality experience for touring virtual card vaults with live price overlays. Walk through your collection in a virtual gallery.'),
        ('Biometric Trading Guard', 'Multi-factor biometric verification (fingerprint, face recognition) for executing high-value transactions. Prevents unauthorized portfolio actions.'),
        ('Mnemonic Propagation', 'Predicts how cultural moments create card demand waves. When a documentary or movie is announced, models the expected impact on related card prices.'),
        ('Circadian Optimizer', 'Analyzes marketplace activity patterns by time of day and day of week. Recommends optimal listing and buying times for maximum value.'),
    ]

    for title, desc in innovations:
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)

    doc.add_page_break()

    # =====================================================================
    # 15. MOBILE & ACCESSIBILITY
    # =====================================================================
    doc.add_heading('15. Mobile & Accessibility', level=1)

    doc.add_heading('15.1 Progressive Web App', level=2)
    doc.add_paragraph(
        'MSI is a full PWA installable on any mobile device. Features include:'
    )
    bullets(doc, [
        'Home screen installation on iOS and Android',
        'Offline access to your collection data',
        'Background sync for price updates when connectivity resumes',
        'Push notifications for watchlist alerts and price movements',
        'Haptic feedback for touch interactions',
        'Swipe gestures for card triage and quick actions',
    ])

    doc.add_heading('15.2 OCR Card Scanner', level=2)
    doc.add_paragraph(
        'Point your phone camera at any card to automatically identify it. OCR reads the '
        'card text (player name, year, set, number) and matches it against the database '
        'for quick cataloging.'
    )

    doc.add_heading('15.3 Command Palette', level=2)
    doc.add_paragraph(
        'Press Ctrl+K (Cmd+K on Mac) from anywhere to open the global command palette. '
        'Type to search for any feature, card, player, or action. Keyboard navigation '
        'support for power users who prefer not to use a mouse.'
    )

    doc.add_heading('15.4 Accessibility', level=2)
    bullets(doc, [
        'ARIA labels throughout the application for screen reader compatibility',
        'Keyboard shortcuts for all major actions',
        'WCAG-compliant color contrast ratios',
        'Responsive layouts from mobile (320px) to ultra-wide desktop',
        'Skeleton loaders for better perceived performance during data fetching',
    ])

    doc.add_page_break()

    # =====================================================================
    # 16. SETTINGS
    # =====================================================================
    doc.add_heading('16. Settings & Account Management', level=1)

    doc.add_heading('16.1 Subscription Management', level=2)
    doc.add_paragraph('Route: /billing')
    doc.add_paragraph(
        'Manage your subscription tier (Free, Pro, Enterprise) via Stripe integration. '
        'View usage metrics for AI features, sync operations, and export operations.'
    )

    doc.add_heading('16.2 Data Sync & Migration', level=2)
    doc.add_paragraph(
        'The platform supports local-first data with cloud sync. The migration engine '
        'handles upgrading from localStorage to Supabase cloud storage, ensuring no data '
        'loss during the transition. Sync scheduler manages background updates.'
    )

    doc.add_heading('16.3 Goal Planner', level=2)
    doc.add_paragraph(
        'Set financial targets for your collection: total value goals, ROI targets, set '
        'completion objectives, and acquisition budgets. Visual progress indicators track '
        'your journey toward each goal.'
    )

    doc.add_heading('16.4 Portfolio Builder', level=2)
    doc.add_paragraph('Route: /builder')
    doc.add_paragraph(
        'Drag-and-drop portfolio organization. Group cards into custom categories, set '
        'allocation targets, and visually manage your portfolio structure.'
    )

    # ---- Footer ----
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('--- End of Feature Guide ---')
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    r.font.size = Pt(10)

    doc.save('/home/user/ModernSportsIntelligenceDemo/Modern_Sports_Intelligence_Feature_Guide.docx')
    print('Feature Guide saved.')


if __name__ == '__main__':
    main()
